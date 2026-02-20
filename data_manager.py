# data_manager.py
import os
import json
import shutil
import docx

class NovelProject:
    def __init__(self, root_path):
        self.root_path = root_path
        self.meta_path = os.path.join(self.root_path, "meta.json")
        self.meta = {
            "title": os.path.basename(self.root_path),
            "global_synopsis": "",
            "characters": [],
            "volumes": []
        }
        self.load_meta()

    def load_meta(self):
        if os.path.exists(self.meta_path):
            with open(self.meta_path, 'r', encoding='utf-8') as f:
                self.meta = json.load(f)
        else:
            self.save_meta()

    def save_meta(self):
        with open(self.meta_path, 'w', encoding='utf-8') as f:
            json.dump(self.meta, f, ensure_ascii=False, indent=4)

    def add_volume(self, vol_name, synopsis=""):
        vol_path = os.path.join(self.root_path, vol_name)
        if not os.path.exists(vol_path):
            os.makedirs(vol_path)
        self.meta["volumes"].append({"name": vol_name, "synopsis": synopsis, "chapters": []})
        self.save_meta()

    def add_chapter(self, vol_index, chap_name, synopsis="", ai_synopsis=""):
        vol_name = self.meta["volumes"][vol_index]["name"]
        chap_filename = f"{chap_name}.docx"
        chap_path = os.path.join(self.root_path, vol_name, chap_filename)

        if not os.path.exists(chap_path):
            doc = docx.Document()
            doc.save(chap_path)

        self.meta["volumes"][vol_index]["chapters"].append({
            "name": chap_name,
            "synopsis": synopsis,
            "ai_synopsis": ai_synopsis
        })
        self.save_meta()

    def rename_volume(self, v_idx, new_name):
        old_name = self.meta["volumes"][v_idx]["name"]
        if old_name == new_name: return
        old_path = os.path.join(self.root_path, old_name)
        new_path = os.path.join(self.root_path, new_name)
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
        self.meta["volumes"][v_idx]["name"] = new_name
        self.save_meta()

    def rename_chapter(self, v_idx, c_idx, new_name):
        vol_name = self.meta["volumes"][v_idx]["name"]
        old_name = self.meta["volumes"][v_idx]["chapters"][c_idx]["name"]
        if old_name == new_name: return
        old_path = os.path.join(self.root_path, vol_name, f"{old_name}.docx")
        new_path = os.path.join(self.root_path, vol_name, f"{new_name}.docx")
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
        self.meta["volumes"][v_idx]["chapters"][c_idx]["name"] = new_name
        self.save_meta()

    def delete_volume(self, v_idx):
        vol_name = self.meta["volumes"][v_idx]["name"]
        vol_path = os.path.join(self.root_path, vol_name)
        if os.path.exists(vol_path):
            shutil.rmtree(vol_path)
        del self.meta["volumes"][v_idx]
        self.save_meta()

    def delete_chapter(self, v_idx, c_idx):
        vol_name = self.meta["volumes"][v_idx]["name"]
        chap_name = self.meta["volumes"][v_idx]["chapters"][c_idx]["name"]
        chap_path = os.path.join(self.root_path, vol_name, f"{chap_name}.docx")
        if os.path.exists(chap_path):
            os.remove(chap_path)
        del self.meta["volumes"][v_idx]["chapters"][c_idx]
        self.save_meta()

    def read_chapter_content(self, vol_name, chap_name):
        chap_path = os.path.join(self.root_path, vol_name, f"{chap_name}.docx")
        if os.path.exists(chap_path):
            doc = docx.Document(chap_path)
            return "\n".join([p.text for p in doc.paragraphs])
        return ""

    def save_chapter_content(self, vol_name, chap_name, content):
        chap_path = os.path.join(self.root_path, vol_name, f"{chap_name}.docx")
        doc = docx.Document()
        for line in content.split('\n'):
            doc.add_paragraph(line)
        doc.save(chap_path)