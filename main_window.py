# main_window.py
import os
import docx
from PyQt6.QtWidgets import (QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel,
                             QTextEdit, QPushButton, QScrollArea, QSplitter, QMessageBox,
                             QFileDialog, QTreeWidget, QTreeWidgetItem, QMenu, QStackedWidget,
                             QInputDialog, QToolBar, QCheckBox)
from PyQt6.QtCore import Qt, QSettings
from PyQt6.QtGui import QShortcut, QKeySequence, QAction, QTextDocument
from PyQt6.QtPrintSupport import QPrinter
from data_manager import NovelProject
from ai_worker import AutoPilotWorker,AIWorker,CorrectionWorker
from ui_components import SettingsDialog, CharacterWidget
from PyQt6.QtWidgets import QToolButton, QMenu, QListWidget, QDockWidget # æ–°å¢å¼•ç”¨

class MainWindow(QMainWindow):
    def __init__(self, project_path):
        super().__init__()
        self.project = NovelProject(project_path)
        self.settings = QSettings("AIWriter", "Settings")
        self.character_widgets = []
        self.current_vol_index = -1
        self.current_chap_index = -1
        self.switch_project = False
        self.is_generating = False

        self.gen_v_idx = -1  # æ­£åœ¨ç”Ÿæˆçš„å·ç´¢å¼•
        self.gen_c_idx = -1  # æ­£åœ¨ç”Ÿæˆçš„ç« ç´¢å¼•
        self.gen_content_buffer = ""  # æ­£æ–‡ç”Ÿæˆçš„å†…å­˜ç¼“å†²åŒº
        self.gen_reasoning_buffer = ""  # æ€è€ƒè¿‡ç¨‹çš„å†…å­˜ç¼“å†²åŒº

        self.setWindowTitle(f"AI ç½‘æ–‡è¾…åŠ©åˆ›ä½œç³»ç»Ÿ - ğŸ“– [{self.project.meta['title']}] (æŒ‰ Ctrl+S ä¿å­˜)")
        self.resize(1400, 850)

        self.init_menu_and_toolbar()
        self.init_ui()
        self.setup_shortcuts()
        self.refresh_tree()

    def init_menu_and_toolbar(self):
        # èœå•æ 
        menubar = self.menuBar()
        file_menu = menubar.addMenu('æ–‡ä»¶')

        settings_action = QAction('âš™ï¸ å…¨å±€/å¤§æ¨¡å‹è®¾ç½®', self)
        settings_action.triggered.connect(self.open_settings)
        file_menu.addAction(settings_action)

        # æ˜¾çœ¼çš„é¡¶éƒ¨å·¥å…·æ  (ä»»ä½•æ—¶å€™éƒ½å¯ä»¥å¿«é€Ÿè°ƒå‡ºè®¾ç½®)
        toolbar = QToolBar("Main Toolbar")
        toolbar.setMovable(False)
        self.addToolBar(toolbar)

        # ====== è¿”å›é¦–é¡µæŒ‰é’® ======
        btn_home = QPushButton("ğŸ  è¿”å›é¦–é¡µ")
        btn_home.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #E6A23C;")
        btn_home.clicked.connect(self.return_to_home)
        toolbar.addWidget(btn_home)

        # ====== ä¸€é”®æˆä¹¦æŒ‰é’® ======
        btn_export = QPushButton("ğŸ“š ä¸€é”®æˆä¹¦")
        btn_export.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #67C23A;")
        btn_export.clicked.connect(self.export_book)
        toolbar.addWidget(btn_export)

        toolbar.addSeparator()

        btn_settings = QPushButton("âš™ï¸ è®¾ç½®æ¨¡å‹å‚æ•°")
        btn_settings.setStyleSheet("background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold;")
        btn_settings.clicked.connect(self.open_settings)
        toolbar.addWidget(btn_settings)

        toolbar.addSeparator()

        lbl_status = QLabel("  ğŸ’¡ æç¤ºï¼šåœ¨å·¦ä¾§æ ‘çŠ¶å›¾å³é”®å¯æ–°å»ºå·/ç« ã€‚")
        lbl_status.setStyleSheet("color: #909399; font-size: 13px;")
        toolbar.addWidget(lbl_status)

        # ====== ã€ä¿®å¤ã€‘æ‰¾å›ä¸¢å¤±çš„è‡ªåŠ¨æŒ‚æœºæŒ‰é’® ======
        self.btn_auto_pilot = QPushButton("ğŸ¤– å¼€å¯è‡ªåŠ¨æŒ‚æœº")
        self.btn_auto_pilot.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #9C27B0;")
        self.btn_auto_pilot.clicked.connect(self.toggle_auto_pilot)
        toolbar.addWidget(self.btn_auto_pilot)

        # ====== å…¨æ–‡ä¸€é”®çº é”™èœå• ======
        self.btn_full_correct = QToolButton()
        self.btn_full_correct.setText("ğŸ©º å…¨æ–‡ä¸€é”®çº é”™")
        self.btn_full_correct.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self.btn_full_correct.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #F56C6C; padding: 5px;")

        self.full_correct_menu = QMenu(self)
        self.full_correct_menu.addAction("é”™åˆ«å­—/è¯­ç—…çº é”™", lambda: self.start_correction("full", "typo"))
        self.full_correct_menu.addAction("è®¾å®š/é€»è¾‘çº é”™", lambda: self.start_correction("full", "setting"))
        self.full_correct_menu.addAction("ğŸŒŸ å…¨éƒ¨çº é”™", lambda: self.start_correction("full", "all"))
        self.btn_full_correct.setMenu(self.full_correct_menu)

        toolbar.addWidget(self.btn_full_correct)

        # ====== ä¾§è¾¹æ å¼€å…³ ======
        self.btn_toggle_log = QPushButton("ğŸ“‹ çº é”™æ—¥å¿—")
        self.btn_toggle_log.setStyleSheet("background-color: transparent; border: none; color: #909399;")
        self.btn_toggle_log.setCheckable(True)
        self.btn_toggle_log.clicked.connect(self.toggle_log_sidebar)
        toolbar.addWidget(self.btn_toggle_log)

    def open_settings(self):
        SettingsDialog(self).exec()

    def init_ui(self):
        central_widget = QWidget()
        central_widget.setStyleSheet("background-color: #FFFFFF;")  # è®©ä¸»å†…å®¹åŒºä¿æŒç™½è‰²æ¸…çˆ½
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(10, 10, 10, 10)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)

        # ====== å·¦ä¾§ï¼šæ–‡ä»¶æ ‘å¯¼èˆª ======
        tree_container = QWidget()
        tree_layout = QVBoxLayout(tree_container)
        tree_layout.setContentsMargins(0, 0, 0, 0)

        self.tree = QTreeWidget()
        self.tree.setHeaderLabel("ğŸ“š å°è¯´å¤§çº²ç›®å½• (å³é”®æ“ä½œ)")
        self.tree.header().setStyleSheet("font-weight: bold; font-size: 15px; color: #303133;")
        self.tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.tree.customContextMenuRequested.connect(self.show_context_menu)
        self.tree.itemClicked.connect(self.on_tree_select)

        tree_layout.addWidget(self.tree)

        # ====== ä¸­é—´ï¼šè®¾å®šé¢æ¿åŒº (Stacked) ======
        self.stacked_widget = QStackedWidget()

        # é¡µé¢0: å…¨å±€è®¾å®š
        self.page_global = QWidget()
        gl_layout = QVBoxLayout(self.page_global)
        gl_layout.setContentsMargins(10, 0, 10, 0)
        gl_layout.addWidget(QLabel("<span style='font-size:16px; font-weight:bold;'>ğŸŒ å…¨å±€æ•…äº‹æ¢—æ¦‚</span>"))
        self.story_synopsis_input = QTextEdit(self.project.meta["global_synopsis"])
        gl_layout.addWidget(self.story_synopsis_input)

        gl_layout.addWidget(
            QLabel("<span style='font-size:16px; font-weight:bold; margin-top:10px;'>ğŸ‘¥ æ ¸å¿ƒäººç‰©è®¾å®š</span>"))
        self.char_list_layout = QVBoxLayout()
        scroll_char = QScrollArea()
        scroll_char.setWidgetResizable(True)
        scroll_char.setStyleSheet("border: none;")
        char_container = QWidget()
        char_container.setLayout(self.char_list_layout)
        scroll_char.setWidget(char_container)
        gl_layout.addWidget(scroll_char)

        btn_add_char = QPushButton("â• æ·»åŠ æ–°äººç‰©")
        btn_add_char.setStyleSheet("border-style: dashed; background-color: #FAFAFA;")
        btn_add_char.clicked.connect(self.add_character)
        gl_layout.addWidget(btn_add_char)

        btn_save_global = QPushButton("ğŸ’¾ ä¿å­˜å…¨å±€è®¾å®š")
        btn_save_global.setStyleSheet("background-color: #409EFF; color: white; font-weight: bold; border: none;")
        btn_save_global.clicked.connect(self.save_global_meta)
        gl_layout.addWidget(btn_save_global)

        # é¡µé¢1: å·è®¾å®š
        self.page_volume = QWidget()
        vl_layout = QVBoxLayout(self.page_volume)
        self.lbl_vol_title = QLabel("<b>å½“å‰å·: </b>")
        self.lbl_vol_title.setStyleSheet("font-size: 16px; color: #303133;")
        vl_layout.addWidget(self.lbl_vol_title)
        self.vol_synopsis_input = QTextEdit()
        self.vol_synopsis_input.setPlaceholderText("æœ¬å·çš„æ ¸å¿ƒä¸»çº¿ã€å‰§æƒ…èµ°å‘æ¢—æ¦‚...")
        vl_layout.addWidget(self.vol_synopsis_input)
        btn_save_vol = QPushButton("ğŸ’¾ ä¿å­˜å·è®¾å®š")
        btn_save_vol.setStyleSheet("background-color: #409EFF; color: white; font-weight: bold; border: none;")
        btn_save_vol.clicked.connect(self.save_vol_meta)
        vl_layout.addWidget(btn_save_vol)

        # é¡µé¢2: ç« è®¾å®š
        self.page_chapter = QWidget()
        cl_layout = QVBoxLayout(self.page_chapter)
        self.lbl_chap_title = QLabel("<b>å½“å‰ç« : </b>")
        self.lbl_chap_title.setStyleSheet("font-size: 16px; color: #303133;")
        cl_layout.addWidget(self.lbl_chap_title)
        self.chap_synopsis_input = QTextEdit()
        self.chap_synopsis_input.setPlaceholderText("æœ¬ç« ç»†çº²ã€å‡ºåœºäººç‰©ã€ååœºé¢è¦æ±‚...")
        cl_layout.addWidget(self.chap_synopsis_input)
        btn_save_chap = QPushButton("ğŸ’¾ ä¿å­˜ç« è®¾å®š")
        btn_save_chap.setStyleSheet("background-color: #409EFF; color: white; font-weight: bold; border: none;")
        btn_save_chap.clicked.connect(self.save_chap_meta)
        cl_layout.addWidget(btn_save_chap)

        self.stacked_widget.addWidget(self.page_global)
        self.stacked_widget.addWidget(self.page_volume)
        self.stacked_widget.addWidget(self.page_chapter)

        # ====== å³ä¾§ï¼šå†™ä½œè¾“å‡ºåŒº ======
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(10, 0, 0, 0)

        # 1. é¡¶éƒ¨æ“ä½œæŒ‰é’®è¡Œ (æ”¾åœ¨åŒä¸€ä¸ªæ°´å¹³å¸ƒå±€é‡Œ)
        btn_action_layout = QHBoxLayout()

        self.btn_start = QPushButton("ğŸš€ æ ¹æ®è®¾å®šæ’°å†™æœ¬ç« ")
        self.btn_start.setEnabled(False)  # å¿…é¡»é€‰ä¸­ç« èŠ‚æ‰èƒ½å†™
        self.btn_start.setStyleSheet(
            "font-size: 16px; font-weight: bold; background-color: #A0CFFF; color: white; border: none; padding: 12px; border-radius: 6px;"
        )
        self.btn_start.clicked.connect(self.start_generation)
        btn_action_layout.addWidget(self.btn_start)

        # ç« èŠ‚ä¸€é”®çº é”™æŒ‰é’®
        self.btn_chap_correct = QToolButton()
        self.btn_chap_correct.setText("ğŸ”§ ç« èŠ‚ä¸€é”®çº é”™")
        self.btn_chap_correct.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self.btn_chap_correct.setStyleSheet(
            "font-size: 14px; font-weight: bold; background-color: #E6A23C; color: white; border: none; padding: 10px; border-radius: 6px;")
        self.btn_chap_correct.setEnabled(False)  # åˆå§‹åŒ–æ—¶æœªé€‰ç« èŠ‚ä¸å¯ç”¨

        self.chap_correct_menu = QMenu(self)
        self.chap_correct_menu.addAction("é”™åˆ«å­—/è¯­ç—…çº é”™", lambda: self.start_correction("chapter", "typo"))
        self.chap_correct_menu.addAction("è®¾å®š/é€»è¾‘çº é”™", lambda: self.start_correction("chapter", "setting"))
        self.chap_correct_menu.addAction("ğŸŒŸ å…¨éƒ¨çº é”™", lambda: self.start_correction("chapter", "all"))
        self.btn_chap_correct.setMenu(self.chap_correct_menu)
        btn_action_layout.addWidget(self.btn_chap_correct)

        # å°†æŒ‰é’®è¡ŒåŠ å…¥å³ä¾§ä¸»å¸ƒå±€
        right_layout.addLayout(btn_action_layout)

        # 2. æ€è€ƒè¿‡ç¨‹æ˜¾ç¤ºåŒº
        self.btn_toggle_thinking = QPushButton("ğŸ”½ æ”¶èµ·æ€è€ƒè¿‡ç¨‹")
        self.btn_toggle_thinking.setStyleSheet(
            "background-color: transparent; border: none; color: #909399; text-align: left;")
        self.btn_toggle_thinking.clicked.connect(self.toggle_thinking)
        right_layout.addWidget(self.btn_toggle_thinking)

        self.thinking_output = QTextEdit()
        self.thinking_output.setReadOnly(True)
        self.thinking_output.setStyleSheet(
            "background-color: #F8F9FA; color: #8A8F99; border: 1px solid #E4E7ED; border-radius: 6px;")
        self.thinking_output.setFixedHeight(120)
        right_layout.addWidget(self.thinking_output)

        # 3. æ­£æ–‡æ˜¾ç¤ºåŒº
        right_layout.addWidget(
            QLabel("<span style='font-size:16px; font-weight:bold;'>âœï¸ å°è¯´æ­£æ–‡åŒº (æŒ‰ Ctrl+S å®æ—¶ä¿å­˜åˆ° docx)</span>"))

        self.content_output = QTextEdit()
        self.content_output.setStyleSheet("""
                    font-size: 16px; 
                    line-height: 1.8; 
                    padding: 15px; 
                    color: #2C3E50;
                    background-color: #FAFAFA;
                """)
        right_layout.addWidget(self.content_output)

        # ====== ã€æ–°å¢ã€‘æ—¥å¿—ä¾§è¾¹æ  (ä½œä¸ºä¸€ä¸ªå¯éšè—çš„ QListWidget) ======
        self.log_list = QListWidget()
        self.log_list.setStyleSheet(
            "background-color: #FAFAFA; border: 1px solid #E4E7ED; color: #606266; padding: 5px;")
        self.log_list.setWordWrap(True)  # ã€æ–°å¢ã€‘å¼€å¯è‡ªåŠ¨æ¢è¡Œï¼Œé˜²æ­¢æ—¥å¿—è¿‡é•¿éš¾ä»¥æŸ¥çœ‹
        self.log_list.hide()  # é»˜è®¤éšè—

        # å°†åŸæœ‰çš„ right_widget åŒ…è£…è¿›å¦ä¸€ä¸ª Splitterï¼Œä½¿å…¶èƒ½å’Œæ—¥å¿—æ å·¦å³æ‹–æ‹½
        right_splitter = QSplitter(Qt.Orientation.Horizontal)
        right_splitter.addWidget(right_widget)
        right_splitter.addWidget(self.log_list)
        right_splitter.setSizes([800, 200])  # è®¾å®šåˆå§‹æ¯”ä¾‹

        # æœ€åæŠŠæ‰€æœ‰å¤§æ¿å—æ‹¼è£…è¿›æœ€å¤–å±‚çš„ splitter
        splitter.addWidget(tree_container)
        splitter.addWidget(self.stacked_widget)
        splitter.addWidget(right_splitter)  # è¿™é‡Œç”¨æ–°çš„ right_splitter æ›¿æ¢æ‰äº†å•çº¯çš„ right_widget
        splitter.setSizes([250, 300, 850])

        # åˆå§‹åŒ–åŠ è½½äººç‰©
        for char_data in self.project.meta.get("characters", []):
            self.add_character(char_data)

    def setup_shortcuts(self):
        shortcut_save = QShortcut(QKeySequence("Ctrl+S"), self)
        shortcut_save.activated.connect(self.save_all)
        shortcut_delete = QShortcut(QKeySequence("Delete"), self.tree)
        shortcut_delete.activated.connect(lambda: self.ui_delete_item(self.tree.currentItem()))

    # --- UI è¾…åŠ©ä¸äº¤äº’é€»è¾‘ ---
    def add_character(self, init_data=None):
        widget = CharacterWidget(self.remove_character, init_data)
        self.char_list_layout.addWidget(widget)
        self.character_widgets.append(widget)

    def remove_character(self, widget):
        self.char_list_layout.removeWidget(widget)
        widget.deleteLater()
        self.character_widgets.remove(widget)

    def toggle_thinking(self):
        is_visible = self.thinking_output.isVisible()
        self.thinking_output.setVisible(not is_visible)
        self.btn_toggle_thinking.setText("ğŸ”½ æ”¶èµ·æ€è€ƒè¿‡ç¨‹" if not is_visible else "â–¶ï¸ å±•å¼€æ€è€ƒè¿‡ç¨‹")

    # === æ—¥å¿—ä¾§è¾¹æ åˆ‡æ¢ ===
    def toggle_log_sidebar(self, checked):
        if checked:
            self.log_list.show()
        else:
            self.log_list.hide()

    def update_ui_state(self):
        # 1. æ£€æŸ¥å½“å‰è§†è§’çš„ç« èŠ‚æ˜¯å¦æ­£åœ¨è¢«å¤§æ¨¡å‹æ’°å†™ã€çº é”™æˆ–æ­£åœ¨è‡ªåŠ¨æŒ‚æœº
        is_auto_piloting = getattr(self, 'is_auto_piloting', False)
        is_correcting = getattr(self, 'is_correcting', False)
        is_generating = getattr(self, 'is_generating', False)

        is_viewing_gen_chapter = ((is_generating or is_auto_piloting) and
                                  self.current_vol_index == getattr(self, 'gen_v_idx', -1) and
                                  self.current_chap_index == getattr(self, 'gen_c_idx', -1))

        # åªè¦æ­£åœ¨ç”Ÿæˆå½“å‰ç« ï¼Œæˆ–è€…å¤„äºå…¨å±€æŒ‚æœºã€å…¨æ–‡/å•ç« çº é”™çŠ¶æ€ï¼Œä¸¥æ ¼é”å®šæ–‡æœ¬æ¡†ä¸ºåªè¯»
        self.content_output.setReadOnly(is_viewing_gen_chapter or is_auto_piloting or is_correcting)

        # 2. åŠ¨æ€æ”¹å˜ç”ŸæˆæŒ‰é’®çš„é¢œè‰²å’Œæ–‡æ¡ˆ
        if is_generating:
            self.btn_start.setEnabled(True)
            if is_viewing_gen_chapter:
                self.btn_start.setText("ğŸ›‘ åœæ­¢ç”Ÿæˆ (æ­£åœ¨è¾“å‡ºå½“å‰ç« )")
                self.btn_start.setStyleSheet(
                    "font-size: 15px; font-weight: bold; background-color: #F56C6C; color: white; border: none; padding: 12px; border-radius: 6px;")
            else:
                self.btn_start.setText("ğŸ›‘ åœæ­¢åå°ç”Ÿæˆ (å…¶ä»–ç« æ­£åœ¨ç å­—)")
                self.btn_start.setStyleSheet(
                    "font-size: 15px; font-weight: bold; background-color: #E6A23C; color: white; border: none; padding: 12px; border-radius: 6px;")
        elif is_auto_piloting:
            # ã€é‡è¦ã€‘ç¡®ä¿æŒ‚æœºæ—¶å•ç« ç”ŸæˆæŒ‰é’®ä¾ç„¶æ˜¯ç¦ç”¨çŠ¶æ€
            self.btn_start.setEnabled(False)
            self.btn_start.setText("ğŸ¤– æŒ‚æœºæ¨¡å¼è¿›è¡Œä¸­...")
            self.btn_start.setStyleSheet(
                "font-size: 16px; font-weight: bold; background-color: #A0CFFF; color: white; border: none; padding: 12px; border-radius: 6px;")
        else:
            if self.current_chap_index != -1:
                self.btn_start.setText("ğŸš€ æ ¹æ®è®¾å®šæ’°å†™æœ¬ç« ")
                self.btn_start.setEnabled(True)
                self.btn_start.setStyleSheet(
                    "font-size: 16px; font-weight: bold; background-color: #67C23A; color: white; border: none; padding: 12px; border-radius: 6px;")
            else:
                self.btn_start.setText("ğŸš€ æ ¹æ®è®¾å®šæ’°å†™æœ¬ç« ")
                self.btn_start.setEnabled(False)
                self.btn_start.setStyleSheet(
                    "font-size: 16px; font-weight: bold; background-color: #A0CFFF; color: white; border: none; padding: 12px; border-radius: 6px;")

        # ====== çº é”™æŒ‰é’®çŠ¶æ€æ›´æ–° ====== (ä¿ç•™ä¸‹åŠéƒ¨åˆ†çš„è¿™ä¸€è¡Œï¼Œä¸åšä¿®æ”¹)
        has_chap_selected = (self.current_chap_index != -1)
        is_correcting = getattr(self, 'is_correcting', False)

        if is_correcting:
            # çº é”™è¿›è¡Œä¸­ï¼šæŠŠå…¨æ–‡çº é”™æŒ‰é’®é­”æ”¹æˆçº¢è‰²çš„â€œåœæ­¢â€æŒ‰é’®
            self.btn_full_correct.setEnabled(True)
            self.btn_full_correct.setText("ğŸ›‘ åœæ­¢çº é”™")
            self.btn_full_correct.setMenu(None)  # æ‘˜æ‰ä¸‹æ‹‰èœå•ï¼Œå˜æˆæ™®é€šå¯ç‚¹å‡»æŒ‰é’®
            self.btn_full_correct.setStyleSheet(
                "background-color: #F56C6C; border: 1px solid #DCDFE6; font-weight:bold; color: white; padding: 5px;"
            )
            # ç»‘å®šåœæ­¢äº‹ä»¶ï¼ˆå…ˆé™é»˜è§£ç»‘ä»¥é˜²é‡å¤ç»‘å®šï¼‰
            try:
                self.btn_full_correct.clicked.disconnect()
            except Exception:
                pass
            self.btn_full_correct.clicked.connect(self.cancel_correction)

            # æŠŠå¦ä¸€ä¸ªçº é”™æŒ‰é’®å½»åº•é”æ­»
            self.btn_chap_correct.setEnabled(False)
            self.btn_chap_correct.setText("ğŸ”§ çº é”™è¿è¡Œä¸­...")
        else:
            # æ¢å¤æ­£å¸¸ï¼šæŠŠèœå•è£…å›å»ï¼Œæ”¹å›åŸæ¥çš„é¢œè‰²
            self.btn_full_correct.setEnabled(True)
            self.btn_full_correct.setText("ğŸ©º å…¨æ–‡ä¸€é”®çº é”™")
            self.btn_full_correct.setMenu(self.full_correct_menu)
            self.btn_full_correct.setStyleSheet(
                "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #F56C6C; padding: 5px;"
            )
            # è§£é™¤ç‚¹å‡»äº‹ä»¶çš„ç»‘å®šï¼Œæ¢å¤èœå•å±•å¼€é»˜è®¤è¡Œä¸º
            try:
                self.btn_full_correct.clicked.disconnect()
            except Exception:
                pass

            self.btn_chap_correct.setEnabled(has_chap_selected)
            self.btn_chap_correct.setText("ğŸ”§ ç« èŠ‚ä¸€é”®çº é”™")
            self.btn_chap_correct.setMenu(self.chap_correct_menu)
            self.btn_chap_correct.setStyleSheet(
                "font-size: 14px; font-weight: bold; background-color: #E6A23C; color: white; border: none; padding: 10px; border-radius: 6px;"
            )

    def start_correction(self, scope, mode):
        api_key = self.settings.value("api_key", "")
        if not api_key:
            QMessageBox.warning(self, "é”™è¯¯", "ç¼ºå°‘ API Keyï¼")
            return

        if scope == "chapter" and (self.current_vol_index == -1 or self.current_chap_index == -1):
            QMessageBox.warning(self, "é”™è¯¯", "è¯·å…ˆåœ¨å·¦ä¾§é€‰æ‹©è¦çº é”™çš„ç« èŠ‚ï¼")
            return

        warning_msg = "å…¨æ–‡çº é”™å°†æ‰«æå…¨ä¹¦æ‘˜è¦ï¼Œå¯èƒ½æ¶ˆè€—å¤§é‡ Token ä¸”è€—æ—¶è¾ƒé•¿ï¼Œè¯·è€å¿ƒç­‰å¾…ã€‚ç¡®è®¤å¼€å§‹å—ï¼Ÿ" if scope == "full" else "å³å°†ä½¿ç”¨ AI æ£€æŸ¥å¹¶ä¿®æ”¹å½“å‰ç« èŠ‚åŸæ–‡ï¼Œç¡®è®¤å¼€å§‹å—ï¼Ÿ"
        if QMessageBox.question(self, 'å¯åŠ¨çº é”™', warning_msg) != QMessageBox.StandardButton.Yes:
            return

        self.save_all()  # å¼ºåˆ¶ä¿å­˜å½“å‰æœ€æ–°çŠ¶æ€
        self.is_correcting = True
        self.update_ui_state()

        # è‡ªåŠ¨å±•å¼€å¹¶æ¸…ç©ºä¾§è¾¹æ å‡†å¤‡è®°å½•
        self.btn_toggle_log.setChecked(True)
        self.toggle_log_sidebar(True)
        self.log_list.addItem(f"=== å¼€å§‹æ–°çš„çº é”™ä»»åŠ¡ ({'å…¨ä¹¦' if scope == 'full' else 'å•ç« '}) ===")
        self.log_list.scrollToBottom()

        base_url = self.settings.value("base_url", "https://api.deepseek.com")
        model = self.settings.value("model", "deepseek-reasoner")
        temp = float(self.settings.value("temperature", 0.7))

        self.correct_worker = CorrectionWorker(api_key, base_url, model, temp, self.project, scope, mode)
        if scope == "chapter":
            self.correct_worker.set_target(self.current_vol_index, self.current_chap_index)

        self.correct_worker.status_signal.connect(lambda msg: self.statusBar().showMessage(msg))
        self.correct_worker.log_signal.connect(self.append_correction_log)
        self.correct_worker.update_text_signal.connect(self.apply_corrected_text)
        self.correct_worker.finished_signal.connect(self.correction_finished)
        self.correct_worker.error_signal.connect(self.handle_error)

        # ====== ã€æ–°å¢ã€‘è¿æ¥æ€è€ƒè¿‡ç¨‹ä¿¡å·ï¼Œå¹¶åœ¨å¯åŠ¨æ—¶æ¸…ç©ºä¸”å±•å¼€æ€è€ƒé¢æ¿ ======
        self.correct_worker.reasoning_signal.connect(self.append_thinking)
        self.thinking_output.clear()
        if not self.thinking_output.isVisible():
            self.toggle_thinking()

        self.correct_worker.start()

    def append_correction_log(self, log_msg):
        self.log_list.addItem(log_msg)
        self.log_list.scrollToBottom()

    def apply_corrected_text(self, v_idx, c_idx, new_content, new_summary):
        vol_name = self.project.meta["volumes"][v_idx]["name"]
        chap = self.project.meta["volumes"][v_idx]["chapters"][c_idx]
        chap_name = chap["name"]

        # åå°è½ç›˜
        self.project.save_chapter_content(vol_name, chap_name, new_content)
        if new_summary and new_summary != chap.get("ai_synopsis", ""):
            chap["ai_synopsis"] = new_summary
            self.project.save_meta()

        # å¦‚æœå½“å‰ UI æ­£å¥½åœç•™åœ¨è¢«ä¿®æ”¹çš„è¿™ä¸€ç« ï¼Œå®æ—¶åˆ·æ–°æ–‡æœ¬æ¡†
        if self.current_vol_index == v_idx and self.current_chap_index == c_idx:
            self.content_output.setText(new_content)
            self.statusBar().showMessage(f"âœ¨ å½“å‰ç« èŠ‚ [{chap_name}] çº é”™å¹¶åˆ·æ–°å®Œæ¯•ï¼", 3000)

    def correction_finished(self):
        self.is_correcting = False
        self.update_ui_state()  # è¿™ä¸€æ­¥ä¼šè®©æŒ‰é’®ä»â€œåœæ­¢â€é‡æ–°å˜å›â€œä¸€é”®çº é”™â€

        # åˆ¤æ–­æ˜¯è¢«æ‰‹åŠ¨åœæ­¢çš„è¿˜æ˜¯è‡ªç„¶è·‘å®Œçš„
        if hasattr(self, 'correct_worker') and self.correct_worker._is_cancelled:
            self.statusBar().showMessage("ğŸ›‘ çº é”™ä»»åŠ¡å·²æ‰‹åŠ¨ç»ˆæ­¢ï¼", 3000)
            self.log_list.addItem("=== çº é”™å·²æ‰‹åŠ¨ç»ˆæ­¢ ===")
        else:
            self.statusBar().showMessage("âœ… çº é”™ä»»åŠ¡å…¨éƒ¨å®Œæˆï¼", 3000)
            self.log_list.addItem("=== çº é”™ä»»åŠ¡ç»“æŸ ===")

        self.log_list.scrollToBottom()

    # --- ç›®å½•æ ‘é€»è¾‘ ---
    def refresh_tree(self):
        self.tree.clear()
        root = QTreeWidgetItem(self.tree, [self.project.meta["title"]])
        root.setIcon(0, self.style().standardIcon(self.style().StandardPixmap.SP_DirIcon))
        root.setData(0, Qt.ItemDataRole.UserRole, {"type": "root"})

        for v_idx, vol in enumerate(self.project.meta["volumes"]):
            v_node = QTreeWidgetItem(root, [vol["name"]])
            v_node.setIcon(0, self.style().standardIcon(self.style().StandardPixmap.SP_FileDialogDetailedView))
            v_node.setData(0, Qt.ItemDataRole.UserRole, {"type": "volume", "v_idx": v_idx})

            for c_idx, chap in enumerate(vol["chapters"]):
                c_node = QTreeWidgetItem(v_node, [chap["name"]])
                c_node.setIcon(0, self.style().standardIcon(self.style().StandardPixmap.SP_FileIcon))
                c_node.setData(0, Qt.ItemDataRole.UserRole, {"type": "chapter", "v_idx": v_idx, "c_idx": c_idx})
        self.tree.expandAll()

    def show_context_menu(self, position):
        item = self.tree.itemAt(position)
        menu = QMenu()
        menu.setStyleSheet(
            "QMenu { background-color: white; border: 1px solid #DCDFE6; } QMenu::item:selected { background-color: #ECF5FF; color: #409EFF; }")

        if not item or item.data(0, Qt.ItemDataRole.UserRole)["type"] == "root":
            action_add_vol = menu.addAction("ğŸ“ æ–°å»ºå·")
            action_add_vol.triggered.connect(self.ui_add_volume)
        elif item.data(0, Qt.ItemDataRole.UserRole)["type"] == "volume":
            action_add_chap = menu.addAction("ğŸ“„ åœ¨æ­¤å·ä¸‹æ–°å»ºç« ")
            v_idx = item.data(0, Qt.ItemDataRole.UserRole)["v_idx"]
            action_add_chap.triggered.connect(lambda: self.ui_add_chapter(v_idx))
            # ã€æ–°å¢ã€‘å·çš„ä¿®æ”¹ä¸åˆ é™¤
            action_rename = menu.addAction("âœï¸ é‡å‘½åå·")
            action_rename.triggered.connect(lambda: self.ui_rename_item(item))
            action_delete = menu.addAction("ğŸ—‘ï¸ åˆ é™¤å·")
            action_delete.triggered.connect(lambda: self.ui_delete_item(item))

        elif item.data(0, Qt.ItemDataRole.UserRole)["type"] == "chapter":
            # ã€æ–°å¢ã€‘ç« çš„ä¿®æ”¹ä¸åˆ é™¤
            action_rename = menu.addAction("âœï¸ é‡å‘½åç« ")
            action_rename.triggered.connect(lambda: self.ui_rename_item(item))
            action_delete = menu.addAction("ğŸ—‘ï¸ åˆ é™¤ç« ")
            action_delete.triggered.connect(lambda: self.ui_delete_item(item))

        menu.exec(self.tree.viewport().mapToGlobal(position))

    def ui_add_volume(self):
        text, ok = QInputDialog.getText(self, "æ–°å»ºå·", "è¯·è¾“å…¥å·å:")
        if ok and text:
            self.project.add_volume(text)
            self.refresh_tree()

    def ui_add_chapter(self, v_idx):
        text, ok = QInputDialog.getText(self, "æ–°å»ºç« ", "è¯·è¾“å…¥ç« å:")
        if ok and text:
            self.project.add_chapter(v_idx, text)
            self.refresh_tree()

    def ui_rename_item(self, item):
        if not item: return
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if data["type"] == "root": return

        old_name = item.text(0)
        item_type = "å·" if data["type"] == "volume" else "ç« "

        new_name, ok = QInputDialog.getText(self, f"é‡å‘½å{item_type}", f"è¯·è¾“å…¥æ–°çš„{item_type}å:", text=old_name)
        if ok and new_name and new_name.strip() != old_name:
            new_name = new_name.strip()
            # æ‰§è¡Œæ•°æ®é‡å‘½å
            if data["type"] == "volume":
                self.project.rename_volume(data["v_idx"], new_name)
            elif data["type"] == "chapter":
                self.project.rename_chapter(data["v_idx"], data["c_idx"], new_name)

            # åˆ·æ–°æ ‘ä¸å³ä¾§æ ‡é¢˜æ˜¾ç¤º
            self.refresh_tree()
            if data["type"] == "volume" and self.current_vol_index == data["v_idx"]:
                self.lbl_vol_title.setText(f"<b>å½“å‰å·: {new_name}</b>")
            elif data["type"] == "chapter" and self.current_vol_index == data["v_idx"] and self.current_chap_index == \
                    data["c_idx"]:
                vol_name = self.project.meta["volumes"][data["v_idx"]]["name"]
                self.lbl_chap_title.setText(f"<b>å½“å‰ç« : {vol_name} - {new_name}</b>")

    def ui_delete_item(self, item):
        if not item: return
        data = item.data(0, Qt.ItemDataRole.UserRole)
        if data["type"] == "root": return

        if self.is_generating:
            if (data["type"] == "volume" and data["v_idx"] == self.gen_v_idx) or \
                    (data["type"] == "chapter" and data["v_idx"] == self.gen_v_idx and data["c_idx"] == self.gen_c_idx):
                QMessageBox.warning(self, "æ“ä½œå—é™", "è¯¥å·/ç« æ­£åœ¨åå°ç–¯ç‹‚ç å­—ä¸­ï¼Œè¯·å…ˆåœæ­¢ç”Ÿæˆåå†å°è¯•åˆ é™¤ï¼")
                return

        item_type = "å·" if data["type"] == "volume" else "ç« "
        item_name = item.text(0)

        # è¯»å–ç”¨æˆ·æ˜¯å¦å¼€å¯äº†â€œåˆ é™¤å‰ç¡®è®¤â€è®¾ç½®
        needs_confirm = self.settings.value("confirm_delete", True, type=bool)

        if needs_confirm:
            msgBox = QMessageBox(self)
            msgBox.setWindowTitle("ç¡®è®¤åˆ é™¤")
            msgBox.setText(f"ç¡®å®šè¦åˆ é™¤{item_type}ã€{item_name}ã€‘å—ï¼Ÿ\nåˆ é™¤æ“ä½œåŒæ—¶ä¼šç§»é™¤æœ¬åœ°æ–‡ä»¶ï¼Œä¸”ä¸å¯æ¢å¤ï¼")
            msgBox.setIcon(QMessageBox.Icon.Warning)
            msgBox.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            msgBox.setDefaultButton(QMessageBox.StandardButton.No)

            # æ¤å…¥â€œä¸å†æé†’â€çš„ CheckBox
            cb = QCheckBox("ä»¥åä¸å†æé†’")
            msgBox.setCheckBox(cb)

            if msgBox.exec() != QMessageBox.StandardButton.Yes:
                return  # ç”¨æˆ·å–æ¶ˆäº†åˆ é™¤

            # å¦‚æœå‹¾é€‰äº†ä¸å†æé†’ï¼Œæ›´æ–°å…¨å±€è®¾ç½®
            if cb.isChecked():
                self.settings.setValue("confirm_delete", False)

        # æ‰§è¡Œåˆ é™¤
        if data["type"] == "volume":
            self.project.delete_volume(data["v_idx"])
        elif data["type"] == "chapter":
            self.project.delete_chapter(data["v_idx"], data["c_idx"])

        # åˆ é™¤åï¼Œé‡ç½®å³ä¾§ç¼–è¾‘é¢æ¿å›åˆ°å…¨å±€è®¾å®šé¡µ
        self.stacked_widget.setCurrentIndex(0)
        self.current_vol_index = -1
        self.current_chap_index = -1
        self.update_ui_state()
        self.refresh_tree()

    def on_tree_select(self, item):
        # ã€æ–°å¢ã€‘åœ¨åˆ‡æ¢ç›®å½•ä¹‹å‰ï¼Œå…ˆé™é»˜ä¿å­˜å½“å‰é€‰ä¸­çš„å·/ç« ä¿¡æ¯ï¼Œé˜²æ­¢å†…å®¹ä¸¢å¤±
        self.save_all(silent=True)

        data = item.data(0, Qt.ItemDataRole.UserRole)
        self.current_vol_index = -1
        self.current_chap_index = -1

        if data["type"] == "root":
            self.stacked_widget.setCurrentIndex(0)

        elif data["type"] == "volume":
            v_idx = data["v_idx"]
            self.current_vol_index = v_idx
            vol_data = self.project.meta["volumes"][v_idx]
            self.lbl_vol_title.setText(f"<b>å½“å‰å·: {vol_data['name']}</b>")
            self.vol_synopsis_input.setText(vol_data.get("synopsis", ""))
            self.stacked_widget.setCurrentIndex(1)

        elif data["type"] == "chapter":
            v_idx = data["v_idx"]
            c_idx = data["c_idx"]
            self.current_vol_index = v_idx
            self.current_chap_index = c_idx

            vol_data = self.project.meta["volumes"][v_idx]
            chap_data = vol_data["chapters"][c_idx]

            self.lbl_chap_title.setText(f"<b>å½“å‰ç« : {vol_data['name']} - {chap_data['name']}</b>")
            self.chap_synopsis_input.setText(chap_data.get("synopsis", ""))
            self.stacked_widget.setCurrentIndex(2)

            # ã€å…³é”®ä¿®å¤ã€‘ï¼šå¢åŠ å¯¹æŒ‚æœºçŠ¶æ€ is_auto_piloting çš„åˆ¤æ–­ï¼Œå¦åˆ™æŒ‚æœºæ—¶ä¼šè¢«å½“ä½œæ™®é€šæŸ¥çœ‹ï¼Œå¯¼è‡´æ€è€ƒè¿‡ç¨‹è¢« clear()
            is_active_gen = getattr(self, 'is_generating', False) or getattr(self, 'is_auto_piloting', False)

            if is_active_gen and getattr(self, 'gen_v_idx', -1) == v_idx and getattr(self, 'gen_c_idx', -1) == c_idx:
                # å¦‚æœåˆ‡å›äº†æ­£åœ¨ç”Ÿæˆçš„ç« ï¼Œå±•ç¤ºå†…å­˜ä¸­çš„å®æ—¶æµ
                self.content_output.setText(self.gen_content_buffer)
                self.thinking_output.setText(self.gen_reasoning_buffer)
                # æ»šåŠ¨æ¡ç§»åˆ°æœ€åº•ç«¯
                self.content_output.moveCursor(self.content_output.textCursor().MoveOperation.End)
                self.thinking_output.moveCursor(self.thinking_output.textCursor().MoveOperation.End)
            else:
                # æŸ¥çœ‹å…¶ä»–ç« èŠ‚ï¼Œè¯»å–æœ¬åœ°è®°å½•
                content = self.project.read_chapter_content(vol_data["name"], chap_data["name"])
                self.content_output.setText(content)
                self.thinking_output.clear()
        self.update_ui_state()

    def return_to_home(self):
        reply = QMessageBox.question(self, 'è¿”å›é¦–é¡µ', 'ç¡®å®šè¦é€€å‡ºå½“å‰é¡¹ç›®å¹¶è¿”å›é¦–é¡µå—ï¼Ÿ\n(ç³»ç»Ÿå°†è‡ªåŠ¨ä¿å­˜å½“å‰è¿›åº¦)',
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.save_all()  # è‡ªåŠ¨ä¿å­˜å½“å‰æ•°æ®
            self.switch_project = True  # è®¾ç½®æ ‡å¿—ä½ä¸º True
            self.close()  # å…³é—­å½“å‰ä¸»çª—å£

    # --- æ•°æ®ä¿å­˜é€»è¾‘ ---
    def save_global_meta(self, silent=False):
        self.project.meta["global_synopsis"] = self.story_synopsis_input.toPlainText().strip()
        chars = []
        for w in self.character_widgets:
            d = w.get_data()
            if any(d.values()):
                chars.append(d)
        self.project.meta["characters"] = chars
        self.project.save_meta()
        if not silent:
            QMessageBox.information(self, "æç¤º", "å…¨å±€è®¾å®šä¿å­˜æˆåŠŸï¼")

    def save_vol_meta(self, silent=False):
        if self.current_vol_index != -1:
            self.project.meta["volumes"][self.current_vol_index][
                "synopsis"] = self.vol_synopsis_input.toPlainText().strip()
            self.project.save_meta()
            if not silent:
                QMessageBox.information(self, "æç¤º", "å½“å‰å·è®¾å®šä¿å­˜æˆåŠŸï¼")

    def save_chap_meta(self, silent=False):
        if self.current_chap_index != -1:
            self.project.meta["volumes"][self.current_vol_index]["chapters"][self.current_chap_index][
                "synopsis"] = self.chap_synopsis_input.toPlainText().strip()
            self.project.save_meta()
            if not silent:
                QMessageBox.information(self, "æç¤º", "å½“å‰ç« è®¾å®šä¿å­˜æˆåŠŸï¼")

    def save_all(self, silent=True):
        # 1. å¦‚æœåœ¨å…¨å±€é¡µï¼Œä¿å­˜å…¨å±€ï¼›å¦‚æœåœ¨å·é¡µï¼Œä¿å­˜å·ï¼›å¦‚æœåœ¨ç« é¡µï¼Œä¿å­˜ç« æ¢—æ¦‚å’Œæ­£æ–‡
        idx = self.stacked_widget.currentIndex()
        if idx == 0:
            self.save_global_meta(silent=silent)
        elif idx == 1:
            self.save_vol_meta(silent=silent)
        elif idx == 2:
            self.save_chap_meta(silent=silent)
            # ä¿å­˜ docx æ­£æ–‡
            if self.current_vol_index != -1 and self.current_chap_index != -1:
                vol_name = self.project.meta["volumes"][self.current_vol_index]["name"]
                chap_name = self.project.meta["volumes"][self.current_vol_index]["chapters"][self.current_chap_index][
                    "name"]
                self.project.save_chapter_content(vol_name, chap_name, self.content_output.toPlainText())

        # æ— è®ºæ˜¯å¿«æ·é”®è¿˜æ˜¯åˆ‡æ¢ç« èŠ‚è§¦å‘ï¼Œéƒ½åœ¨åº•éƒ¨çŠ¶æ€æ æä¾›æ— æ„Ÿæç¤º
        self.statusBar().showMessage("âœ… å°è¯´æ­£æ–‡åŠè®¾å®šå·²è‡ªåŠ¨ä¿å­˜ï¼", 3000)

    def export_book(self):
        # 1. å¼ºåˆ¶ä¿å­˜å½“å‰æœ€æ–°è¿›åº¦
        self.save_all()

        # 2. å¼¹å‡ºä¿å­˜æ–‡ä»¶å¯¹è¯æ¡†
        file_path, filter_type = QFileDialog.getSaveFileName(
            self,
            "ä¸€é”®æˆä¹¦ - é€‰æ‹©å¯¼å‡ºä½ç½®",
            f"{self.project.meta['title']}.docx",
            "Word æ–‡æ¡£ (*.docx);;Markdown æ–‡æ¡£ (*.md);;çº¯æ–‡æœ¬ (*.txt);;PDF æ–‡æ¡£ (*.pdf)"
        )

        if not file_path:
            return

        # 3. æ ¹æ®åç¼€åè°ƒç”¨ç›¸åº”çš„å¯¼å‡ºæ–¹æ³•
        try:
            ext = os.path.splitext(file_path)[1].lower()
            title = self.project.meta['title']

            if ext == '.docx':
                self._export_docx(file_path, title)
            elif ext == '.md':
                self._export_md(file_path, title)
            elif ext == '.txt':
                self._export_txt(file_path, title)
            elif ext == '.pdf':
                self._export_pdf(file_path, title)

            QMessageBox.information(self, "å¯¼å‡ºæˆåŠŸ", f"æ­å–œï¼å°è¯´å·²æˆåŠŸå¯¼å‡ºè‡³ï¼š\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "å¯¼å‡ºå¤±è´¥", f"å¯¼å‡ºè¿‡ç¨‹ä¸­å‘ç”Ÿé”™è¯¯ï¼š\n{str(e)}")

    def _export_docx(self, file_path, title):
        doc = docx.Document()
        doc.add_heading(title, 0)  # ä¹¦åä½œä¸ºä¸»æ ‡é¢˜

        for vol in self.project.meta["volumes"]:
            doc.add_heading(vol["name"], level=1)  # å·åä½œä¸ºä¸€çº§æ ‡é¢˜
            for chap in vol["chapters"]:
                doc.add_heading(chap["name"], level=2)  # ç« åä½œä¸ºäºŒçº§æ ‡é¢˜
                content = self.project.read_chapter_content(vol["name"], chap["name"])
                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
        doc.save(file_path)

    def _export_txt(self, file_path, title):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"ã€Š{title}ã€‹\n\n")
            for vol in self.project.meta["volumes"]:
                f.write(f"ã€{vol['name']}ã€‘\n\n")
                for chap in vol["chapters"]:
                    f.write(f"  {chap['name']}\n\n")
                    content = self.project.read_chapter_content(vol["name"], chap["name"])
                    f.write(f"{content}\n\n")

    def _export_md(self, file_path, title):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"# {title}\n\n")
            for vol in self.project.meta["volumes"]:
                f.write(f"## {vol['name']}\n\n")
                for chap in vol["chapters"]:
                    f.write(f"### {chap['name']}\n\n")
                    content = self.project.read_chapter_content(vol["name"], chap["name"])
                    f.write(f"{content}\n\n")

    def _export_pdf(self, file_path, title):
        # PDF å¯¼å‡ºåˆ©ç”¨ PyQt6 è‡ªå¸¦çš„å¯Œæ–‡æœ¬è½¬æ¢ä¸º HTML å†æ¸²æŸ“æ‰“å°çš„æœºåˆ¶
        html_content = f"<h1 style='text-align: center;'>{title}</h1>"
        for vol in self.project.meta["volumes"]:
            html_content += f"<h2 style='color: #2C3E50;'>{vol['name']}</h2>"
            for chap in vol["chapters"]:
                html_content += f"<h3>{chap['name']}</h3>"
                content = self.project.read_chapter_content(vol["name"], chap["name"])
                for line in content.split('\n'):
                    if line.strip():
                        html_content += f"<p style='text-indent: 2em; line-height: 1.5;'>{line.strip()}</p>"

        document = QTextDocument()
        document.setHtml(html_content)

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)

        # æ¸²æŸ“ç”Ÿæˆ PDF
        document.print(printer)
    # --- æ ¸å¿ƒå¤§æ¨¡å‹ç”Ÿæˆé€»è¾‘ (åŒ…å«å¤æ‚çš„ä¸Šä¸‹æ–‡ç»„è£…) ---
    def build_prompts(self):
        meta = self.project.meta

        # 1. ç»„è£…å…¨å±€è®¾å®š
        global_story = meta.get("global_synopsis", "æœªæä¾›ã€‚")
        char_texts = [f"ã€{c['name']}ã€‘ æ€§åˆ«:{c['gender']} æ€§æ ¼:{c['personality']} ç»å†:{c['experience']}" for c in
                      meta.get("characters", [])]
        char_setting = "\n".join(char_texts) if char_texts else "æœªæä¾›æ˜ç¡®äººç‰©ã€‚"

        system_prompt = f"""ä½ æ˜¯ä¸€ä½ç»éªŒä¸°å¯Œçš„ç½‘æ–‡å¤§ç¥ä½œå®¶ã€‚è¯·æ ¹æ®å…¨å±€è®¾å®šå’Œä¸Šä¸‹æ–‡è¿è´¯åœ°æ’°å†™å°è¯´æ­£æ–‡ã€‚

    ã€å…¨å±€æ•…äº‹å¤§çº²ã€‘
    {global_story}

    ã€æ ¸å¿ƒäººç‰©è®¾å®šã€‘
    {char_setting}

    ã€å†™ä½œè¦æ±‚ã€‘
    1. ä¸¥æ ¼éµå¾ªä¸–ç•Œè§‚ã€äººè®¾å’Œå‰§æƒ…é€»è¾‘ã€‚
    2. åŠ¨ä½œã€ç¥æ€ã€å¿ƒç†æå†™ç”ŸåŠ¨ï¼Œç¬¦åˆç½‘æ–‡çˆ½æ„ŸèŠ‚å¥ã€‚
    3. ç›´æ¥è¾“å‡ºæ­£æ–‡ï¼Œç¦æ­¢ä»»ä½•è§£é‡Šæ€§åºŸè¯å’Œå¤šä½™çš„å¯’æš„ã€‚
    4. ã€é‡è¦ã€‘åœ¨æ­£æ–‡è¾“å‡ºå®Œæ¯•åï¼Œå¿…é¡»å¦èµ·ä¸€è¡Œå¹¶ä¸¥æ ¼ä»¥ `[AI_SUMMARY]` ä½œä¸ºåˆ†å‰²ç¬¦ï¼Œç„¶åè¾“å‡ºçº¦300å­—çš„æœ¬ç« è¯¦ç»†æ¢—æ¦‚ï¼ˆå¿…é¡»åŒ…å«å…·ä½“å‘ç”Ÿçš„æƒ…èŠ‚ã€äººç‰©å‘å±•ã€æ–°å‡ºç°çš„ç‰©å“/äººç‰©ä»¥åŠåŸ‹ä¸‹çš„ä¼ç¬”ï¼‰ã€‚æ­¤éƒ¨åˆ†ä»…ç”¨äºç³»ç»Ÿå†…éƒ¨è®°å½•ã€‚"""

        # 2. ç»„è£…å†å²ä¸Šä¸‹æ–‡ä¸ä¸Šä¸€ç« å†…å®¹
        past_context = ""
        prev_chapter_content = ""

        v_idx = self.current_vol_index
        c_idx = self.current_chap_index

        prev_v_idx, prev_c_idx = -1, -1
        if c_idx > 0:
            prev_v_idx, prev_c_idx = v_idx, c_idx - 1
        elif v_idx > 0:
            for i in range(v_idx - 1, -1, -1):
                if len(meta["volumes"][i]["chapters"]) > 0:
                    prev_v_idx = i
                    prev_c_idx = len(meta["volumes"][i]["chapters"]) - 1
                    break

        if prev_v_idx != -1 and prev_c_idx != -1:
            pv_name = meta["volumes"][prev_v_idx]["name"]
            pc_name = meta["volumes"][prev_v_idx]["chapters"][prev_c_idx]["name"]
            prev_chapter_content = self.project.read_chapter_content(pv_name, pc_name)
            if len(prev_chapter_content) > 1500:
                prev_chapter_content = "...(å‰æ–‡çœç•¥)...\n" + prev_chapter_content[-1500:]

        # ã€ä¿®æ”¹å¤„ã€‘æå–è¿‡å¾€æ‰€æœ‰æ¢—æ¦‚æ—¶ï¼Œä¼˜å…ˆä½¿ç”¨ ai_synopsis
        history_str = ""
        for i in range(v_idx + 1):
            vol = meta["volumes"][i]
            history_str += f"\n> {vol['name']} (æ¢—æ¦‚: {vol.get('synopsis', 'æ— ')})\n"

            chap_limit = c_idx if i == v_idx else len(vol["chapters"])
            for j in range(chap_limit):
                chap = vol["chapters"][j]

                # ä¼˜å…ˆè¯»å– AI ä¹‹å‰ç”Ÿæˆçš„æ¢—æ¦‚ï¼Œå¦‚æœæ²¡æœ‰åˆ™é™çº§è¯»å–ç”¨æˆ·çš„ç»†çº²
                ai_syn = chap.get("ai_synopsis", "")
                user_syn = chap.get("synopsis", "æ— ")
                display_syn = ai_syn if ai_syn.strip() else user_syn

                history_str += f"  - {chap['name']}: {display_syn}\n"

        if len(history_str) > 15000:
            history_str = "ã€æ³¨æ„ï¼šå› å‰æ–‡è¿‡é•¿ï¼Œæ­¤å¤„ä»…æä¾›è¿‡å¾€å·æ¢—æ¦‚ã€‘\n"
            for i in range(v_idx + 1):
                vol = meta["volumes"][i]
                history_str += f"\n> {vol['name']} (æ¢—æ¦‚: {vol.get('synopsis', 'æ— ')})\n"

        if not history_str.strip():
            history_str = "æœ¬ä¹¦åˆšåˆšå¼€ç¯‡ï¼Œæ— è¿‡å¾€å†å²ã€‚"

        curr_vol = meta["volumes"][v_idx]
        curr_chap = curr_vol["chapters"][c_idx]

        user_prompt = f"""è¯·ä¸ºæˆ‘æ’°å†™æœ€æ–°ç« èŠ‚çš„æ­£æ–‡ã€‚

ã€è¿‡å¾€å‰§æƒ…è½¨è¿¹å‚è€ƒã€‘
{history_str.strip()}

"""
        if prev_chapter_content.strip():
            user_prompt += f"ã€ç´§æ¥ä¸Šä¸€ç« çš„æœ«å°¾å†…å®¹ã€‘(è¯·ä¿è¯å‰§æƒ…å’Œå¯¹è¯çš„è¿è´¯è¿‡æ¸¡)\n{prev_chapter_content.strip()}\n\n"

        user_prompt += f"""ã€æœ¬æ¬¡å†™ä½œä»»åŠ¡ã€‘
å½“å‰æ‰€å¤„å·ï¼š{curr_vol['name']}
æœ¬å·æ ¸å¿ƒæ¢—æ¦‚ï¼š{curr_vol.get('synopsis', 'æ— ')}

å½“å‰éœ€æ’°å†™ç« èŠ‚ï¼š{curr_chap['name']}
æœ¬ç« ç»†çº²è¦æ±‚ï¼š{curr_chap.get('synopsis', 'æ— ')}

ã€è¡ŒåŠ¨æŒ‡ä»¤ã€‘
è¯·æ ¹æ®æœ¬ç« ç»†çº²è¦æ±‚ï¼Œé¡ºç€ä¸Šä¸€ç« çš„æƒ…èŠ‚å±•å¼€ï¼Œæ‰©å†™ä¸ºæ–‡ç¬”æµç•…çš„å®Œæ•´æ­£æ–‡ï¼
ã€é‡è¦ã€‘åœ¨æ­£æ–‡è¾“å‡ºå®Œæ¯•åï¼Œå¿…é¡»å¦èµ·ä¸€è¡Œå¹¶ä¸¥æ ¼ä»¥ `[AI_SUMMARY]` ä½œä¸ºåˆ†å‰²ç¬¦ï¼Œç„¶åè¾“å‡ºçº¦300å­—é«˜åº¦ç»“æ„åŒ–çš„ã€æœ¬ç« å¤ç›˜ä¸è®°å¿†é”šç‚¹ã€‘
    åœ¨ `[AI_SUMMARY]` ä¹‹åï¼Œå¿…é¡»ä¸¥æ ¼æŒ‰ç…§ä»¥ä¸‹3ä¸ªç»´åº¦è¾“å‡ºï¼ˆå®¢è§‚ã€ç²¾ç‚¼ï¼Œçº¯ä½œå†…éƒ¨è®°å¿†ä½¿ç”¨ï¼‰ï¼š
    1. æ ¸å¿ƒå‰§æƒ…è„‰ç»œï¼šæŒ‰æ—¶é—´é¡ºåºç®€è¿°æœ¬ç« å‘ç”Ÿçš„å®è´¨æ€§äº‹ä»¶ï¼ˆèµ·å› ã€ç»è¿‡ã€ç»“æœï¼‰ã€‚
    2. äººç‰©çŠ¶æ€æ›´æ–°ï¼šè®°å½•æœ¬ç« ä¸»è§’åŠé…è§’çš„è¡Œä¸ºåŠå¿ƒæ€ã€‚
    3. ç‰©å“è®¾å®šæ›´æ–°ï¼šè®°å½•æœ¬ç« æ‰€æœ‰ç‰©å“çŠ¶æ€

"""

        return system_prompt, user_prompt

    def start_generation(self):
        if getattr(self, 'is_generating', False):
            if hasattr(self, 'worker') and self.worker.isRunning():
                self.worker.cancel()
            self.btn_start.setText("ğŸ›‘ æ­£åœ¨åœæ­¢...")
            self.btn_start.setEnabled(False)
            return
        api_key = self.settings.value("api_key", "")
        if not api_key:
            QMessageBox.warning(self, "é”™è¯¯", "ç¼ºå°‘ API Keyï¼Œè¯·ç‚¹å‡»ä¸Šæ–¹ã€âš™ï¸ è®¾ç½®æ¨¡å‹å‚æ•°ã€‘æŒ‰é’®è¿›è¡Œé…ç½®ï¼")
            self.open_settings()
            return

        # ç”Ÿæˆå‰å¼ºåˆ¶ä¿å­˜å½“å‰çš„æ¢—æ¦‚è®¾å®šï¼Œä»¥å…æç¤ºè¯æ²¡ç”¨åˆ°æœ€æ–°å†…å®¹
        self.save_all()
        system_prompt, user_prompt = self.build_prompts()

        # === è®¾ç½®åå°ç”Ÿæˆçš„ç¯å¢ƒå’Œç¼“å†²åŒº ===
        self.is_generating = True
        self.gen_v_idx = self.current_vol_index
        self.gen_c_idx = self.current_chap_index
        self.gen_content_buffer = ""
        self.gen_reasoning_buffer = ""

        self.content_output.clear()
        self.thinking_output.clear()

        self.hit_summary_delimiter = False
        self.content_output.clear()
        self.thinking_output.clear()
        # åˆ·æ–°ç•Œé¢çŠ¶æ€ (æ ‘çŠ¶å›¾ä¸é”å®šï¼Œä»…é”å®šæ­£æ–‡è¾“å…¥æ¡†ï¼ŒæŒ‰é’®å˜çº¢)
        self.update_ui_state()

        base_url = self.settings.value("base_url", "https://api.deepseek.com")
        model = self.settings.value("model", "deepseek-reasoner")
        temperature = float(self.settings.value("temperature", 1.5))
        max_tokens = int(self.settings.value("max_tokens", 6000))

        self.worker = AIWorker(api_key, base_url, model, temperature, max_tokens, system_prompt, user_prompt)
        self.worker.reasoning_signal.connect(self.append_thinking)
        self.worker.content_signal.connect(self.append_content)
        self.worker.error_signal.connect(self.handle_error)
        self.worker.finished_signal.connect(self.generation_finished)
        self.worker.start()

    def append_thinking(self, text):
        self.gen_reasoning_buffer += text  # æ°¸è¿œå†™è¿›åå°ç¼“å†²åŒº
        # åªæœ‰å½“ç”¨æˆ·æ­£åœç•™åœ¨è¯¥ç« æ—¶ï¼Œæ‰å®æ—¶æ¸²æŸ“åœ¨å±å¹•ä¸Š
        if self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx:
            self.thinking_output.insertPlainText(text)
            self.thinking_output.ensureCursorVisible()

    def append_content(self, text):
        # [cite_start]ã€å…³é”®ä¿®å¤ã€‘ï¼šå®æ—¶å°† AI åå‡ºçš„æ–‡å­—æ‹¼æ¥åˆ°åå°ç¼“å†²åŒºä¸­ [cite: 196]
        self.gen_content_buffer += text

        if "[AI_SUMMARY]" in self.gen_content_buffer:
            if not getattr(self, 'hit_summary_delimiter', False):
                self.hit_summary_delimiter = True
                # è§¦å‘åˆ†å‰²ç¬¦æ—¶ï¼Œå°†æ­£æ–‡çš„æœ€åä¸€éƒ¨åˆ†æ¸…ç†å¹²å‡€æ¸²æŸ“åˆ°UIä¸Šï¼Œä¹‹ååœæ­¢æ›´æ–°UIçš„æ­£æ–‡éƒ¨åˆ†
                if self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx:
                    main_content = self.gen_content_buffer.split("[AI_SUMMARY]")[0].strip()
                    self.content_output.setPlainText(main_content)
                    self.content_output.moveCursor(self.content_output.textCursor().MoveOperation.End)
        else:
            # æ­£å¸¸æ¸²æŸ“æ­£æ–‡
            if self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx:
                # ã€æ ¸å¿ƒä¿®å¤ã€‘ï¼šå…ˆå¼ºåˆ¶å°†å…‰æ ‡ç§»åŠ¨åˆ°æ–‡æœ¬æœ€æœ«å°¾ï¼Œå†æ’å…¥æ–‡æœ¬ã€‚é˜²æ­¢é¼ æ ‡ä¹±ç‚¹å¯¼è‡´æ–‡å­—æ’é”™ä½ç½®ï¼
                cursor = self.content_output.textCursor()
                cursor.movePosition(cursor.MoveOperation.End)
                self.content_output.setTextCursor(cursor)

                # åœ¨æœ«å°¾æ’å…¥æœ€æ–°æ–‡æœ¬æµ
                self.content_output.insertPlainText(text)
                self.content_output.ensureCursorVisible()

    def handle_error(self, err_msg):
        QMessageBox.critical(self, "ç”Ÿæˆé”™è¯¯", f"è¯·æ±‚å‘ç”Ÿå¼‚å¸¸ï¼š\n{err_msg}")
        # ã€ä¿®æ”¹å¤„ã€‘æ ¹æ®å½“å‰çš„æ¨¡å¼ï¼Œè°ƒç”¨å¯¹åº”çš„ç»“æŸ/é‡ç½®æ–¹æ³•
        if getattr(self, 'is_auto_piloting', False):
            self.auto_pilot_finished()
        else:
            self.generation_finished()

    def generation_finished(self):
        if self.gen_v_idx != -1 and self.gen_c_idx != -1:
            vol_name = self.project.meta["volumes"][self.gen_v_idx]["name"]
            chap_data = self.project.meta["volumes"][self.gen_v_idx]["chapters"][self.gen_c_idx]
            chap_name = chap_data["name"]

            # ã€æ ¸å¿ƒä¿®æ”¹ã€‘å°†ç¼“å†²åŒºçš„å†…å®¹æ ¹æ®æ ‡è¯†ç¬¦ä¸€åˆ†ä¸ºäºŒ
            parts = self.gen_content_buffer.split("[AI_SUMMARY]")
            main_content = parts[0].strip()
            ai_summary = parts[1].strip() if len(parts) > 1 else ""

            # 1. ä¿å­˜çº¯å‡€çš„æ­£æ–‡åˆ° docx
            self.project.save_chapter_content(vol_name, chap_name, main_content)

            # 2. å¦‚æœæˆåŠŸç”Ÿæˆäº† AI æ€»ç»“ï¼Œå°†å…¶éšå¼ä¿å­˜åˆ° meta å¹¶åœ¨åå°è½ç›˜
            if ai_summary:
                chap_data["ai_synopsis"] = ai_summary
                self.project.save_meta()

            # 3. å¦‚æœç”¨æˆ·è¿˜åœç•™åœ¨è¿™ä¸ªç« èŠ‚ï¼Œç¡®ä¿æ–‡æœ¬æ¡†é‡Œæ˜¾ç¤ºçš„æ˜¯çº¯å‡€çš„ã€æ²¡æœ‰å°¾å·´çš„æ­£æ–‡
            if self.current_vol_index == self.gen_v_idx and self.current_chap_index == self.gen_c_idx:
                self.content_output.setPlainText(main_content)

        # æ¸…é™¤åå°ç”Ÿæˆæ ‡è®°
        self.is_generating = False
        self.gen_v_idx = -1
        self.gen_c_idx = -1

        # åˆ·æ–° UI çŠ¶æ€æ¢å¤åŸè²Œ
        self.update_ui_state()
        self.statusBar().showMessage("âœ… ç« èŠ‚æ­£æ–‡ç”Ÿæˆå®Œæ¯•ï¼ŒAIå†…éƒ¨çº¿ç´¢æ¢—æ¦‚å·²å…¥åº“ï¼", 3000)

    #è¿½åŠ :è‡ªåŠ¨æŒ‚æœºç±»å‡½æ•°
    def toggle_auto_pilot(self):
        if getattr(self, 'is_auto_piloting', False):
            # åœæ­¢æŒ‚æœº
            if hasattr(self, 'auto_worker') and self.auto_worker.isRunning():
                self.auto_worker.cancel()
                self.btn_auto_pilot.setText("ğŸ›‘ æ­£åœ¨åœæ­¢æŒ‚æœº...")
                self.btn_auto_pilot.setEnabled(False)
            else:
                self.auto_pilot_finished()
            return

        # å¼€å¯æŒ‚æœºå‰æ£€æŸ¥
        api_key = self.settings.value("api_key", "")
        if not api_key:
            QMessageBox.warning(self, "é”™è¯¯", "ç¼ºå°‘ API Keyï¼")
            return

        reply = QMessageBox.question(self, 'é«˜èƒ½é¢„è­¦', 'ç¡®å®šå¼€å¯å…¨è‡ªåŠ¨æŒ‚æœºï¼Ÿ\nAIå°†è‡ªåŠ¨æ¶ˆè€—å¤§é‡Tokenè¡¥å…¨æ‰€æœ‰è®¾å®šå’Œæ­£æ–‡ï¼',
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply != QMessageBox.StandardButton.Yes: return

        self.save_all()
        self.is_auto_piloting = True
        self.btn_auto_pilot.setText("ğŸ›‘ åœæ­¢è‡ªåŠ¨æŒ‚æœº")
        self.btn_auto_pilot.setStyleSheet("background-color: #F56C6C; font-weight:bold; color: white;")

        # ç¦ç”¨æ‰‹åŠ¨å•ç« ç”ŸæˆæŒ‰é’®
        self.btn_start.setEnabled(False)
        self.btn_start.setText("æŒ‚æœºæ¨¡å¼è¿›è¡Œä¸­...")

        base_url = self.settings.value("base_url", "https://api.deepseek.com")
        model = self.settings.value("model", "deepseek-reasoner")  # æŒ‚æœºæ¨èç”¨å¼ºåŠ›æ¨¡å‹
        temp = float(self.settings.value("temperature", 0.7))

        self.auto_worker = AutoPilotWorker(api_key, base_url, model, temp, self.project)

        # ä¿¡å·å¯¹æ¥ï¼šçŠ¶æ€åˆ·æ–°
        self.auto_worker.status_signal.connect(lambda msg: self.statusBar().showMessage(msg))
        self.auto_worker.log_signal.connect(lambda msg: self.thinking_output.append(msg))

        # ä¿¡å·å¯¹æ¥ï¼šæµå¼æ­£æ–‡è¾“å‡ºåˆ°å½“å‰ç•Œé¢ï¼ˆå¹¶æ‹¦æˆª [AI_SUMMARY] è¯¦è§ä½ ä¹‹å‰çš„ä»£ç é€»è¾‘ï¼‰
        self.hit_summary_delimiter = False
        self.auto_worker.content_signal.connect(self.append_content)

        # ã€æ–°å¢ã€‘è¿æ¥æ€è€ƒè¿‡ç¨‹ä¿¡å·
        self.auto_worker.reasoning_signal.connect(self.append_thinking)

        # ã€æ–°å¢ã€‘è¿æ¥åˆ‡æ¢ç« èŠ‚ä¿¡å· (å¿…é¡»ç”¨é˜»å¡è¿æ¥ï¼Œç¡®ä¿UIåˆ‡å®Œå†è¾“å‡ºæ–‡å­—)
        self.auto_worker.start_chapter_signal.connect(self.auto_start_chapter,
                                                      Qt.ConnectionType.BlockingQueuedConnection)

        # ä¿¡å·å¯¹æ¥ï¼šåå°æ•°æ®ç»“æ„ä¿®æ”¹
        self.auto_worker.add_volume_signal.connect(self.auto_add_volume, Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.add_chapter_signal.connect(self.auto_add_chapter, Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.save_content_signal.connect(self.auto_save_content, Qt.ConnectionType.BlockingQueuedConnection)

        self.auto_worker.update_chapter_signal.connect(self.auto_update_chapter,
                                                       Qt.ConnectionType.BlockingQueuedConnection)
        self.auto_worker.update_volume_signal.connect(self.auto_update_volume,
                                                      Qt.ConnectionType.BlockingQueuedConnection)

        self.auto_worker.finished_signal.connect(self.auto_pilot_finished)
        self.auto_worker.error_signal.connect(self.handle_error)

        # ã€å…³é”®ä¿®å¤ã€‘ï¼šç»å¯¹ä¸èƒ½åœ¨è¿™é‡Œ clear() æ–‡æœ¬æ¡†ï¼
        # å¦åˆ™ç¨å auto_worker åˆ‡æ¢ç« èŠ‚æ—¶è§¦å‘çš„ save_all ä¼šæŠŠç©ºæ–‡æœ¬æ¡†è¦†ç›–åˆ°å‰ä¸€ç« ï¼
        self.auto_worker.start()

    def auto_update_volume(self, v_idx, synopsis):
        vol = self.project.meta["volumes"][v_idx]
        vol["synopsis"] = synopsis
        self.project.save_meta()

        # å¦‚æœå½“å‰ UI æ­£å¥½åœç•™åœ¨è¿™ä¸€å·çš„è®¾ç½®ç•Œé¢ï¼Œå®æ—¶åˆ·æ–°æ–‡æœ¬æ¡†
        if self.current_vol_index == v_idx and self.stacked_widget.currentIndex() == 1:
            self.vol_synopsis_input.setText(synopsis)
    # --- ä¾› AutoPilotWorker è·¨çº¿ç¨‹è°ƒç”¨çš„ UI å’Œæ•°æ®æ›´æ–°æ§½å‡½æ•° ---
    def auto_update_chapter(self, v_idx, c_idx, ai_synopsis):
        chap = self.project.meta["volumes"][v_idx]["chapters"][c_idx]
        chap["ai_synopsis"] = ai_synopsis

        # æ ¸å¿ƒé€»è¾‘ï¼šå¦‚æœç”¨æˆ·åŸæœ¬å°±æ²¡æœ‰å†™ synopsisï¼Œé‚£å°±æŠŠ AI å†™çš„å¡åˆ°å°é¢ä¸Šï¼›
        # å¦‚æœç”¨æˆ·å†™äº†ï¼Œé‚£å°±ä¿ç•™ç”¨æˆ·å†™çš„ï¼ŒAI çš„æ‰©å†™åªæ”¾åœ¨éšå¼çš„ ai_synopsis é‡Œä¾›å¤§æ¨¡å‹çœ‹
        if not chap.get("synopsis", "").strip():
            chap["synopsis"] = ai_synopsis

        self.project.save_meta()

        # å¦‚æœå½“å‰ UI æ­£å¥½åœç•™åœ¨è¿™ä¸€ç« ï¼Œåˆ·æ–°ä¸€ä¸‹æ–‡æœ¬æ¡†æ˜¾ç¤º
        if self.current_vol_index == v_idx and self.current_chap_index == c_idx:
            self.chap_synopsis_input.setText(chap.get("synopsis", ""))

    def auto_add_volume(self, name, synopsis):
        self.project.add_volume(name, synopsis)
        self.refresh_tree()
        self.tree.scrollToBottom()

    def auto_add_chapter(self, v_idx, name, ai_synopsis):
        # ã€ä¿®æ”¹å¤„ã€‘å°† ai_synopsis åŒæ—¶ä¹Ÿèµ‹å€¼ç»™ synopsis å­—æ®µï¼Œè¿™æ ·å°±èƒ½åœ¨ UI çš„â€œç« è®¾å®šâ€é‡Œçœ‹åˆ°äº†ï¼
        self.project.add_chapter(v_idx, name, synopsis=ai_synopsis, ai_synopsis=ai_synopsis)
        self.refresh_tree()
        self.tree.scrollToBottom()

    def auto_start_chapter(self, v_idx, c_idx):
        self.gen_v_idx = v_idx
        self.gen_c_idx = c_idx
        self.gen_content_buffer = ""
        self.gen_reasoning_buffer = ""
        self.hit_summary_delimiter = False

        # è‡ªåŠ¨é€‰ä¸­å·¦ä¾§æ ‘çŠ¶å›¾å¯¹åº”çš„ç« èŠ‚èŠ‚ç‚¹
        root = self.tree.topLevelItem(0)
        if root and v_idx < root.childCount():
            v_node = root.child(v_idx)
            if c_idx < v_node.childCount():
                c_node = v_node.child(c_idx)
                # é€‰ä¸­æ ‘èŠ‚ç‚¹
                self.tree.setCurrentItem(c_node)
                # è§¦å‘ç‚¹å‡»äº‹ä»¶ï¼Œè®©å³ä¾§é¢æ¿åˆ‡æ¢åˆ°è¯¥ç« çš„ç©ºç™½ç¼–è¾‘çŠ¶æ€
                self.on_tree_select(c_node)

    def auto_save_content(self, v_idx, c_idx, main_content, ai_summary):
        vol_name = self.project.meta["volumes"][v_idx]["name"]
        chap_name = self.project.meta["volumes"][v_idx]["chapters"][c_idx]["name"]

        # ä¿å­˜æœ¬åœ° docx
        self.project.save_chapter_content(vol_name, chap_name, main_content)
        # æ›´æ–° meta ä¸­çš„ AI æ€»ç»“
        if ai_summary:
            self.project.meta["volumes"][v_idx]["chapters"][c_idx]["ai_synopsis"] = ai_summary
            self.project.save_meta()

        # ã€å…³é”®ä¿®å¤ã€‘ï¼šå–æ¶ˆè¿™è¡Œ clear()ï¼Œå°†æ¸…ç†å·¥ä½œäº¤ç»™ on_tree_select å»è‡ªç„¶è¿‡æ¸¡
        self.hit_summary_delimiter = False

    def auto_pilot_finished(self):
        self.is_auto_piloting = False
        self.btn_auto_pilot.setText("ğŸ¤– å¼€å¯è‡ªåŠ¨æŒ‚æœº")
        self.btn_auto_pilot.setEnabled(True)
        self.btn_auto_pilot.setStyleSheet(
            "background-color: transparent; border: 1px solid #DCDFE6; font-weight:bold; color: #9C27B0;")
        self.update_ui_state()  # æ¢å¤åŸæœ‰æŒ‰é’®çŠ¶æ€

    def cancel_correction(self):
        """æ‰‹åŠ¨æš‚åœ/ç»ˆæ­¢çº é”™ä»»åŠ¡"""
        if getattr(self, 'is_correcting', False):
            if hasattr(self, 'correct_worker') and self.correct_worker.isRunning():
                self.correct_worker.cancel()  # è§¦å‘ Worker å†…çš„å–æ¶ˆæ ‡è®°ï¼Œå¹¶å¼ºè¡Œåˆ‡æ–­ç½‘ç»œæµ
                self.log_list.addItem("âš ï¸ æ¥æ”¶åˆ°åœæ­¢æŒ‡ä»¤ï¼Œæ­£åœ¨ç­‰å¾…å½“å‰è¯·æ±‚å®‰å…¨ä¸­æ–­...")
                self.log_list.scrollToBottom()
                self.statusBar().showMessage("ğŸ›‘ æ­£åœ¨åœæ­¢çº é”™...", 3000)